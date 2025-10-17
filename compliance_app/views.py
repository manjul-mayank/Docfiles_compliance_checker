# compliance_app/views.py

import os
import shutil
import uuid
import json
import logging
from pathlib import Path
from typing import Optional, Dict, Any

from django.shortcuts import render
from django.http import JsonResponse, HttpResponse, FileResponse
from django.views.decorators.csrf import csrf_exempt
from django.conf import settings
from django.core.files.uploadedfile import UploadedFile
import docx
import fitz
import language_tool_python
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
# REMOVED: import openai (since we are using Gemini)

# NEW IMPORT: Google Gemini API
from google import genai 
from google.genai.errors import APIError 

import mimetypes
from .forms import DocumentUploadForm

# --- Configuration & Tools ---
REGISTRY: Dict[str, Any] = {}
tool = language_tool_python.LanguageTool('en-US')
logger = logging.getLogger(__name__)

# --- UTILS (File Handlers) ---

def secure_filename(filename: str) -> str:
    """Sanitizes a filename to keep it safe and consistent."""
    # Note: Assumes settings.UPLOAD_DIR is correctly defined as a Path object
    os.makedirs(settings.UPLOAD_DIR, exist_ok=True)
    keep = "".join(c for c in filename if c.isalnum() or c in (" ", ".", "_", "-")).rstrip()
    return keep.replace(" ", "_")

def extract_text_from_docx(path: Path) -> str:
    doc = docx.Document(str(path))
    return "\n\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])

def extract_text_from_pdf(path: Path) -> str:
    text_pages = []
    try:
        with fitz.open(str(path)) as pdf:
            for page in pdf:
                text_pages.append(page.get_text("text"))
        return "\n\n".join(p.strip() for p in text_pages if p.strip())
    except Exception as e:
        logger.error(f"Error extracting text from PDF {path}: {e}")
        return "" 

def extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".docx":
        return extract_text_from_docx(path)
    elif ext == ".pdf":
        return extract_text_from_pdf(path)
    else:
        # Assuming ALLOWED_EXTENSIONS is defined in settings
        allowed_exts = getattr(settings, 'ALLOWED_EXTENSIONS', ['.docx', '.pdf'])
        raise ValueError(f"Unsupported file type: {ext}. Allowed types: {allowed_exts}")

# --- AI + GRAMMAR ---
def grammar_check(text: str) -> dict:
    matches = tool.check(text)
    return {
        "issue_count": len(matches),
        "issues": [
            {"message": m.message, "replacements": m.replacements, "context": m.context}
            for m in matches
        ]
    }

# --- CORRECTED call_ai function for GEMINI (Using dedicated system_instruction) ---
def call_ai(prompt: str, max_tokens=1000) -> dict:
    api_key = getattr(settings, 'GEMINI_API_KEY', None)
    
    if not api_key:
        logger.error("GEMINI_API_KEY not set in Django settings.")
        return {"content": "LLM not configured. Set GEMINI_API_KEY for compliance checks."}
    
    ai_model = getattr(settings, 'AI_MODEL', 'gemini-2.5-flash')
    
    # 🟢 CRITICAL FIX: Define the system instruction as a dedicated variable
    SYSTEM_INSTRUCTION = (
        "You are a professional document compliance evaluator and editor. "
        "Your response must be only the requested output (either an assessment report or modified text) "
        "without conversational filler, preambles, or explanations."
    )
    
    try:
        client = genai.Client(api_key=api_key) 

        resp = client.models.generate_content(
            model=ai_model,
            # 🟢 FIX: Send only the user prompt text in contents
            contents=[prompt], 
            config=genai.types.GenerateContentConfig(
                max_output_tokens=max_tokens, 
                temperature=0.2,
                # 🟢 FIX: Use the dedicated system_instruction parameter
                system_instruction=SYSTEM_INSTRUCTION
            )
        )
        
        if resp.text:
            return {"content": resp.text}
        else:
            # Enhanced logging for debugging empty responses
            logger.warning(f"Gemini returned empty text. Prompt: {prompt[:100]}...")
            return {"content": "Gemini returned an empty response or the content was blocked by safety settings."}
        
    except APIError as e:
        logger.error(f"Gemini API Error: {e}")
        return {"content": f"Gemini API Error: {e}"}
    except Exception as e:
        logger.error(f"Unexpected AI Error: {e}")
        return {"content": f"Unexpected AI Error: {e}"}


def evaluate_compliance(text: str, guidelines: Optional[str] = None) -> dict:
    # Truncate text to 15000 characters to prevent excessive token usage
    safe_text = text[:15000] 
    # 🟢 TEMPORARY DEBUGGING STEP: Check if text is actually extracted
    if not safe_text.strip():
        logger.error("Extracted text is empty or only whitespace. Check file extraction.")
        return {"grammar": {"issue_count": 0, "issues": []}, "ai_report": "Error: Document text could not be extracted.", "summary": "Extraction failure"}
    # 🟢 END DEBUGGING
    
    prompt = f"Assess this text against the following writing guidelines:\n{guidelines or 'Standard English writing rules'}\n\nText:\n{safe_text}"
    ai_result = call_ai(prompt)
    grammar = grammar_check(safe_text)
    
    return {
        "grammar": grammar,
        "ai_report": ai_result.get("content", "AI evaluation failed."),
        "summary": "Combined grammar and AI compliance report"
    }

# --- DOC MODIFICATION ---
def apply_edits_to_docx(original_path: Path, edited_text: str) -> Path:
    # Creates a new DOCX file with the edited text
    doc = docx.Document()
    for para in edited_text.split("\n\n"):
        doc.add_paragraph(para.strip())
    out_path = settings.UPLOAD_DIR / (original_path.stem + "_modified.docx")
    doc.save(str(out_path))
    return out_path

def create_pdf_from_text(text: str, out_path: Path) -> Path:
    # Creates a new PDF file from text
    c = canvas.Canvas(str(out_path), pagesize=letter)
    width, height = letter
    y = height - 72
    
    for line in text.splitlines():
        chunks = [line[i:i + 100] for i in range(0, len(line), 100)] 
        for chunk in chunks:
            if y < 72:
                c.showPage()
                y = height - 72
            c.drawString(72, y, chunk) 
            y -= 14
    c.save()
    return out_path


# ===================================================
#           DJANGO VIEW FUNCTIONS
# ===================================================

# 1. HOME VIEW
def home(request):
    """Renders the main application page."""
    return render(request, 'compliance_app/index.html')


# 2. UPLOAD VIEW (Fixes the recurring AttributeError)
@csrf_exempt
def upload_document(request): 
    """Handles the file upload and saves document metadata to REGISTRY."""
    if request.method != 'POST':
        return JsonResponse({'detail': 'Method not allowed'}, status=405)

    form = DocumentUploadForm(request.POST, request.FILES)
    if form.is_valid():
        uploaded_file: UploadedFile = form.cleaned_data['file']
        guidelines: str = form.cleaned_data['guidelines']
        
        doc_id = str(uuid.uuid4())
        safe_filename = secure_filename(uploaded_file.name)
        
        file_path = settings.UPLOAD_DIR / f"{doc_id}_{safe_filename}"
        
        with open(file_path, 'wb+') as destination:
            for chunk in uploaded_file.chunks():
                destination.write(chunk)

        REGISTRY[doc_id] = {
            "path": str(file_path),
            "filename": safe_filename,
            "guidelines": guidelines,
            "status": "uploaded"
        }
        
        return JsonResponse({
            "doc_id": doc_id, 
            "filename": safe_filename, 
            "detail": "Upload successful"
        })
    else:
        return JsonResponse({'detail': 'Invalid form data or file missing', 'errors': form.errors}, status=400)


# 3. ASSESS VIEW
@csrf_exempt
def assess_document(request, doc_id):
    """Performs compliance and grammar checks on the uploaded document."""
    if request.method != 'POST':
        return JsonResponse({'detail': 'Method not allowed'}, status=405)
        
    rec = REGISTRY.get(doc_id)
    if not rec:
        return JsonResponse({'detail': 'Document not found'}, status=404)
        
    try:
        path = Path(rec["path"])
        text = extract_text(path)
        report = evaluate_compliance(text, rec.get("guidelines"))
        
        rec["status"] = "assessed"
        rec["last_report"] = report
        
        return JsonResponse({"doc_id": doc_id, "filename": rec["filename"], "report": report})
    except Exception as e:
        logger.error(f"Assessment error for {doc_id}: {e}")
        return JsonResponse({'detail': f'Assessment failed: {e}'}, status=500)


# 4. MODIFY VIEW
@csrf_exempt
def modify_document(request):
    """Applies AI-driven modifications based on user instruction."""
    if request.method != 'POST':
        return JsonResponse({'detail': 'Method not allowed'}, status=405)
    
    try:
        data = json.loads(request.body)
        doc_id = data.get('doc_id')
        instruction = data.get('instruction')
    except json.JSONDecodeError:
        return JsonResponse({'detail': 'Invalid JSON body'}, status=400)
        
    rec = REGISTRY.get(doc_id)
    if not rec:
        return JsonResponse({'detail': 'Document not found'}, status=404)
        
    path = Path(rec["path"])
    
    try:
        original_text = extract_text(path)
        
        prompt = f"REWRITE the following text to comply ONLY with the following instruction. Respond only with the modified text, no conversation or explanation:\n\nINSTRUCTION: {instruction}\n\nORIGINAL TEXT:\n{original_text[:15000]}"
        
        edited_text = call_ai(prompt, max_tokens=4096).get("content", "")
        
        if not edited_text or "AI API Error" in edited_text or "Gemini returned an empty response" in edited_text:
            edited_text = original_text 

        ext = path.suffix.lower()
        
        out_filename_base = path.stem + "_modified"
        if ext == ".docx":
            out_path = settings.UPLOAD_DIR / (out_filename_base + ".docx")
            apply_edits_to_docx(path, edited_text) 
        else:
            out_path = settings.UPLOAD_DIR / (out_filename_base + ".pdf")
            create_pdf_from_text(edited_text, out_path)
            
        new_id = str(uuid.uuid4())
        REGISTRY[new_id] = {
            "path": str(out_path), 
            "filename": out_path.name, 
            "status": "modified", 
            "origin": doc_id,
            "original_ext": ext 
        }
        
        return JsonResponse({
            "modified_doc_id": new_id, 
            "download_endpoint": f"/download/{new_id}",
            "filename": out_path.name
        })
    except Exception as e:
        logger.error(f"Modification error for {doc_id}: {e}")
        return JsonResponse({'detail': f'Modification failed: {e}'}, status=500)


# 5. DOWNLOAD VIEW
def download_file(request, doc_id):
    """Serves the file for download."""
    rec = REGISTRY.get(doc_id)
    if not rec:
        return JsonResponse({'detail': 'File not found'}, status=404)
        
    file_path = Path(rec["path"])
    if not file_path.exists():
        return JsonResponse({'detail': 'File not found on server'}, status=404)

    mime_type = mimetypes.guess_type(rec["filename"])[0] or "application/octet-stream"

    try:
        response = FileResponse(open(file_path, 'rb'), content_type=mime_type)
        response['Content-Disposition'] = f'attachment; filename="{rec["filename"]}"'
        return response
    except Exception as e:
        logger.error(f"Download error for {doc_id}: {e}")
        return JsonResponse({'detail': 'Failed to read file for download'}, status=500)